import os
import re
from flask import Flask
import pytesseract
from PIL import Image
from docx import Document
import pandas as pd
import openpyxl

# 1-CONFIGURACIÓN
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

app = Flask(__name__)

carpeta = r"C:\Users\alanr\OneDrive\Documentos\ProyectoE\fotosTEXT"
carpeta_doc = r"C:\Users\alanr\OneDrive\Documentos\ProyectoE\carpeta_documentos"

os.makedirs(carpeta_doc, exist_ok=True)

excel_path = os.path.join(carpeta_doc, "facturas.xlsx")

print("Directorio actual:", os.getcwd())

# 2-FUNCION PARA EXTRAER DATOS DE FACTURA
def extraer_datos_factura(texto):
    datos = {}

    datos["Num_Factura"] = re.search(r"Numero[:\s]*([A-Z0-9-]+)", texto, re.IGNORECASE)
    datos["Fecha"] = re.search(r"Fecha[:\s]*([\d/]+)", texto, re.IGNORECASE)
    datos["N_pedido"] = re.search(r"Número de pedido[:\s]*([A-Z0-9-]+)", texto, re.IGNORECASE)
    datos["Total"] = re.search(r"Total[:\s]*\$?\s*([\d,.]+)", texto, re.IGNORECASE)

    return {
        "NUM_Factura": datos["Num_Factura"].group(1) if datos["Num_Factura"] else "",
        "N_pedido": datos["N_pedido"].group(1) if datos["N_pedido"] else "",
        "Fecha": datos["Fecha"].group(1) if datos["Fecha"] else "",
        "Total": datos["Total"].group(1) if datos["Total"] else "",
    }
    
# 3-PROCESAMIENTO DE IMÁGENES
for i, archivo in enumerate(os.listdir(carpeta), start=1):

    if not archivo.lower().endswith((".jpg", ".jpeg", ".png")):
        continue

    ruta_imagen = os.path.join(carpeta, archivo)

    #-4 OCR
    imagen = Image.open(ruta_imagen).convert("L")
    texto = pytesseract.image_to_string(imagen, lang="spa")

    # 5-LIMPIEZA DE TEXTO
    texto = re.sub(r"[—–-]{2,}", " ", texto)
    texto = re.sub(r"[-]", " \n ", texto)
    texto = re.sub(r"\s{2,}", " ", texto)
    texto = re.sub(r"[°º]", "o", texto)
    texto = re.sub(r"Sofiware As", "Software", texto)
    texto = re.sub(r"EdICÍÓA", "EDICION", texto)
    texto = re.sub(r"6£", "DE", texto)
    texto = texto.replace("\n\n", "\n")

    print(f"\n=== TEXTO ANALIZADO ({i}): {archivo} ===")
    print(texto)

    #6- GUARDAR WORD
    doc = Document()
    doc.add_heading(f"Texto extraído de imagen {archivo}", level=3)
    doc.add_paragraph(texto)

    nombre_doc = os.path.splitext(archivo)[0] + ".docx"
    ruta_doc = os.path.join(carpeta_doc, nombre_doc)
    doc.save(ruta_doc)

    # 7- EXTRAER Y GUARDAR FACTURA EN EXCEL
    if "factura" in texto.lower() and "total" in texto.lower():
        datos_factura = extraer_datos_factura(texto)
        df_nuevo = pd.DataFrame([datos_factura])

        if os.path.exists(excel_path):
            df_existente = pd.read_excel(excel_path)
            df_final = pd.concat([df_existente, df_nuevo], ignore_index=True)
        else:
            df_final = df_nuevo

        df_final.to_excel(excel_path, index=False)

print("\n✔ Proceso finalizado. Word y Excel generados correctamente.")